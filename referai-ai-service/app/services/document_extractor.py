"""Document text extraction service for PDF and DOCX files."""

from __future__ import annotations

import base64
import io
from typing import Literal

from PyPDF2 import PdfReader
from docx import Document

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentExtractor:
    """Extract text from PDF and DOCX files."""

    @staticmethod
    def extract_text(
        file_content_b64: str, file_name: str
    ) -> tuple[bool, str | None, int | None, int | None, str | None, str | None]:
        """
        Extract text from base64 encoded file.

        Returns:
            (success, text, word_count, page_count, method, error)
        """
        try:
            # Decode base64
            file_bytes = base64.b64decode(file_content_b64)
            file_extension = file_name.lower().split(".")[-1]

            if file_extension == "pdf":
                return DocumentExtractor._extract_pdf(file_bytes)
            elif file_extension in ("docx", "doc"):
                return DocumentExtractor._extract_docx(file_bytes)
            else:
                return (
                    False,
                    None,
                    None,
                    None,
                    None,
                    f"Unsupported file type: {file_extension}",
                )

        except Exception as e:
            logger.exception("document.extraction.failed", error=str(e))
            return False, None, None, None, None, f"Extraction failed: {str(e)}"

    @staticmethod
    def _extract_pdf(
        file_bytes: bytes,
    ) -> tuple[bool, str, int, int, Literal["PyPDF2"], None]:
        """Extract text from PDF using PyPDF2."""
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = "\n\n".join(text_parts)
            cleaned_text = DocumentExtractor._clean_text(full_text)
            word_count = len(cleaned_text.split())
            page_count = len(reader.pages)

            logger.info(
                "document.pdf.extracted",
                pages=page_count,
                words=word_count,
            )

            return True, cleaned_text, word_count, page_count, "PyPDF2", None

        except Exception as e:
            logger.exception("document.pdf.failed", error=str(e))
            raise

    @staticmethod
    def _extract_docx(
        file_bytes: bytes,
    ) -> tuple[bool, str, int, int | None, Literal["python-docx"], None]:
        """Extract text from DOCX using python-docx."""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n\n".join(text_parts)
            cleaned_text = DocumentExtractor._clean_text(full_text)
            word_count = len(cleaned_text.split())

            logger.info(
                "document.docx.extracted",
                paragraphs=len(doc.paragraphs),
                words=word_count,
            )

            return True, cleaned_text, word_count, None, "python-docx", None

        except Exception as e:
            logger.exception("document.docx.failed", error=str(e))
            raise

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text by removing extra whitespace and special characters."""
        # Remove multiple spaces
        text = " ".join(text.split())

        # Remove multiple newlines (keep max 2)
        lines = text.split("\n")
        cleaned_lines = []
        empty_count = 0

        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
                empty_count = 0
            else:
                empty_count += 1
                if empty_count <= 1:  # Keep max 1 empty line
                    cleaned_lines.append("")

        return "\n".join(cleaned_lines).strip()
